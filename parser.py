import os
from abc import ABC, abstractmethod
from datetime import datetime

import pandas as pd
from docx import Document
from telethon import TelegramClient


class TelegramClientSingleton:
    _instance = None

    def __new__(cls, session_name, api_id, api_hash):
        if not cls._instance:
            cls._instance = super(TelegramClientSingleton, cls).__new__(cls)
            cls._instance._client = TelegramClient(session_name, api_id, api_hash)
        return cls._instance._client


class TelegramParser(ABC):
    def __init__(self, api_id: int, api_hash: str):
        self._api_id = api_id
        self._api_hash = api_hash
        self._client = TelegramClientSingleton('anon', self._api_id, self._api_hash)

    @abstractmethod
    async def parse_start(self) -> None:
        pass

    @property
    def client(self) -> TelegramClient:
        return self._client


class ParserChannel(TelegramParser):
    def __init__(self, api_id: int, api_hash: str, channels: tuple, my_channel: str):
        super().__init__(api_id, api_hash)
        self.data = []
        self.doc = Document()
        self.channels = channels
        self.my_channel = my_channel

    async def parse_start(self) -> None:
        for channel in self.channels:
            async for message in self.client.iter_messages(channel, limit=5):
                self.data.append({
                    "ID": message.id,
                    "Date": message.date.strftime('%Y-%m-%d %H:%M:%S'),
                    "Text": message.text
                })
                if message.photo:
                    path = await message.download_media()
                    print(path)
                    await self.client.send_file(self.my_channel, path, caption=message.text)
                    os.remove(path)
                    # await self.client.send_message(self.my_channel, message.text)
            df = pd.DataFrame(self.data)
            df.to_csv(f"output_{channel.replace('https://t.me/', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                      index=False)
        # for channel in self.channels:
        #     async for message in self.client.iter_messages(channel, limit=5):
        #         self.doc.add_paragraph(f"ID: {message.id}")
        #         self.doc.add_paragraph(f"Date: {message.date}")
        #         self.doc.add_paragraph(f"Text: {message.text}")
        #
        #         if message.photo:
        #             await message.download_media("image")
        #
        #     self.doc.save(f"output_{channel.replace('https://t.me/', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")


class ChatParser(TelegramParser):
    def __init__(self, api_id, api_hash, chats: tuple):
        super().__init__(api_id, api_hash)
        self.doc = Document()
        self.chats = chats

    async def parse_start(self):
        for chat in self.chats:
            async for username in self.client.iter_participants(chat, limit=5):
                self.doc.add_paragraph(f"username: @{username.username}")
            self.doc.save(f"output_{chat.replace('https://t.me/', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")


api_id = 20188962
api_hash = '002d5308a5bf7ea900ce9fb953f14b6b'
channels = ('https://t.me/RVvoenkor', 'https://t.me/germanylivetv',)
chats = ('https://t.me/rvvoenkor2',)
my_channel = 'https://t.me/coder_it01'
tg_parser_channel = ParserChannel(api_id, api_hash, channels, my_channel)
tg_parser_chat = ChatParser(api_id, api_hash, chats)


async def main():
    await tg_parser_channel.parse_start()
    # await tg_parser_chat.parse()


with tg_parser_channel.client:
    tg_parser_channel.client.loop.run_until_complete(main())

from telethon import TelegramClient, events
from telethon.tl.functions.messages import GetMessages
import pandas as pd
import datetime
import os

class ParserChannel(TelegramParser):
    def __init__(self, api_id: int, api_hash: str, channels: tuple, my_channel: str):
        super().__init__(api_id, api_hash)
        self.data = []
        self.doc = Document()
        self.channels = channels
        self.my_channel = my_channel

    async def parse_start(self) -> None:
        for channel in self.channels:
            messages = []
            async for message in self.client.iter_messages(channel, limit=None):
                messages.append(message)
                if len(messages) >= 5:
                    break

            # Обработка 5 полученных сообщений
            for message in messages:
                self.data.append({
                    "ID": message.id,
                    "Date": message.date.strftime('%Y-%m-%d %H:%M:%S'),
                    "Text": message.text
                })
                if message.photo:
                    path = await message.download_media()
                    print(path)
                    await self.client.send_file(self.my_channel, path, caption=message.text)
                    os.remove(path)

            df = pd.DataFrame(self.data)
            df.to_csv(f"output_{channel.replace('https://t.me/', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                      index=False)

if __name__ == "__main__":
    api_id = 123456
    api_hash = "your_api_hash"
    channels = ("https://t.me/channel1", "https://t.me/channel2")
    my_channel = "https://t.me/my_channel"

    client = TelegramClient("session_name", api_id, api_hash)
    parser = ParserChannel(api_id, api_hash, channels, my_channel)

    with client:
        client.loop.run_until_complete(parser.parse_start())
