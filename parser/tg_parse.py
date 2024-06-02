from telethon import TelegramClient
from telethon.tl.types import Channel, Chat, User
from telethon.tl.functions.contacts import SearchRequest
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import textwrap
from exception_handler import exception_handler


class TelegramParser(ABC):
    def __init__(self, api_id: int, api_hash: str):
        self._api_id = api_id
        self._api_hash = api_hash
        self._client = TelegramClient('anon', api_id, api_hash)

    @property
    def client(self) -> TelegramClient:
        return self._client


class ShowMyChannels(TelegramParser):
    @exception_handler
    async def get_my_channels(self):
        async for dialog in self.client.iter_dialogs():
            print(
                f'Название(канала, чата) - {dialog.name}'
                f'ID - {dialog.id}'
            )


class TelegramParserChannel(TelegramParser):
    def __init__(self, api_id: int, api_hash: str, channel: str = None, my_channel: str = None):
        super().__init__(api_id, api_hash)
        self.channel = channel
        self.my_channel = my_channel

    @exception_handler
    async def send_posts_to_my_channel(self):
        async for message in self.client.iter_messages(self.channel, limit=6):
            if message.text or message.media:
                if message.text and message.media:
                    await self.client.send_message(self.my_channel, message.text, file=message.media)
                elif message.text:
                    await self.client.send_message(self.my_channel, message.text)
                else:
                    await self.client.send_message(self.my_channel, file=message.media)

    @exception_handler
    async def send_posts_to_my_channel_from_id(self):
        entity = await self.client.get_entity(self.channels)
        async for message in self.client.iter_messages(entity, limit=5):
            if message.text or message.media:
                if message.text and message.media:
                    await self.client.send_message(self.my_channel, message.text, file=message.media)
                elif message.text:
                    await self.client.send_message(self.my_channel, message.text)
                else:
                    await self.client.send_message(self.my_channel, file=message.media)


class TelegramParserChat(TelegramParser):
    def __init__(self, api_id, api_hash, chat, chat_id: int = None):
        super().__init__(api_id, api_hash)
        self.chat = chat
        self.chat_id = chat_id

    @exception_handler
    async def get_messages_in_chat(self) -> None:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        header = document.add_heading('Список сообщений', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=1, cols=2)
        header_row = table.rows[0].cells
        header_row[0].text = 'Пользователь'
        header_row[1].text = 'Сообщение'

        async for message in self.client.iter_messages(self.chat_id, limit=1000):
            if message.text and message.sender:
                username = message.sender.username or message.sender.first_name
                username_display = f"https://t.me/{username}" if username else message.sender.first_name

                message_text = message.text.replace('\n', ' ')
                wrapped_text = textwrap.fill(message_text, width=70)

                row_cells = table.add_row().cells
                row_cells[0].text = username_display
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[1].text = wrapped_text
                row_cells[1].paragraphs[0].runs[0].font.italic = True

        document.save('messages.docx')

    @exception_handler
    async def get_messages_in_chat_html(self) -> None:
        html_content = "<html><head><title>Messages</title></head><body>"
        async for message in self.client.iter_messages(self.chat, limit=1000):
            if message.text and message.sender:
                user_link = f"https://t.me/{message.sender.username}"
                html_content += f"<p>Username: <a href='{user_link}'>{message.sender.username}</a></p>"

                if hasattr(message.sender, 'first_name'):
                    html_content += f"<p>Имя: {message.sender.first_name}</p>"
                else:
                    html_content += "<p>Имя отсутствует</п>"

                if hasattr(message.sender, 'last_name'):
                    html_content += f"<p>Фамилия: {message.sender.last_name}</п>"

                html_content += f"<p>Текст комментария: {message.text}</п>"
                html_content += "<hr>"

        html_content += "</body></html>"

        with open('messages.html', 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

    @exception_handler
    async def get_participants_chat(self):
        count: int = 0
        document = Document()
        async for participant in self.client.iter_participants(self.chat, limit=30):
            profile_link = f"https://t.me/{participant.username}"
            document.add_paragraph(f"User ID: {participant.id}")
            document.add_paragraph(f"Username: {profile_link}")
            document.add_paragraph(f"Имя: {participant.first_name}")
            document.add_paragraph(f"Фамилия: {participant.last_name}")
            document.add_paragraph(f"Телефон: {participant.phone}")
            document.add_paragraph(f"Имеет ли аккаунт премиум статус: {participant.premium}")
            document.add_paragraph("-" * 20)
            count += 1
        document.save('users.docx')
        print(f"Спарсено участников: {count}")

    @exception_handler
    async def get_participants_chat_html(self):
        count = 0
        html_content = "<html><head><title>Participants</title></head><body>"
        async for participant in self.client.iter_participants(self.chat, limit=30):
            profile_link = f"https://t.me/{participant.username}"
            html_content += f"<p>User ID: <a href='{profile_link}'>{participant.id}</a></п>"
            html_content += f"<p>Username: <a href='{profile_link}'>{participant.username if participant.username else 'No username'}</а></п>"
            html_content += f"<p>Имя: {participant.first_name}</п>"
            html_content += f"<p>Фамилия: {participant.last_name}</п>"
            html_content += f"<п>Телефон: {participant.phone}</п>"
            html_content += f"<п>Имеет ли аккаунт премиум статус: {participant.premium}</п>"
            html_content += "<hr>"
            count += 1
        html_content += "</body></html>"

        with open('users.html', 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

        print(f"Спарсено участников: {count}")

    @exception_handler
    async def get_participants_not_link_chat(self):
        chat_entity = await self.client.get_entity(self.chat_id)
        with open('data_users.html', 'w', encoding='utf-8') as file:
            file.write('<html><head><title>Participants</title></head><body>')
            file.write('<h1>Список подписчиков</h1>')
            file.write('<table border="1">')
            file.write('<tr><th>ID</th><th>Username</th><th>First Name</th><th>Last Name</th></tr>')

            async for user in self.client.iter_participants(chat_entity, limit=20):
                user_id = user.id
                username = user.username if user.username else 'None'
                first_name = user.first_name if user.first_name else 'None'
                last_name = user.last_name if user.last_name else 'None'

                if user.username:
                    username_link = f'<a href="https://t.me/{user.username}">{user.username}</a>'
                else:
                    username_link = 'None'

                file.write(f'<tr>'
                           f'<td>{user_id}</td>'
                           f'<td>{username_link}</td>'
                           f'<td>{first_name}</td>'
                           f'<td>{last_name}</td>'
                           f'</tr>')

            file.write('</table>')
            file.write('</body></html>')

    @exception_handler
    async def get_participants_not_link_chat_txt(self):
        chat_entity = await self.client.get_entity(self.chat_id)
        with open('data_users.txt', 'w', encoding='utf-8') as file:
            file.write(f"{'ID':<12} {'Username':<20} {'First Name':<20} {'Last Name':<20}\n")
            file.write("=" * 75 + "\n")
            async for user in self.client.iter_participants(chat_entity, limit=10):
                user_id = user.id
                username = user.username if user.username else 'None'
                first_name = user.first_name if user.first_name else 'None'
                last_name = user.last_name if user.last_name else 'None'
                file.write(f"{user_id:<12} {username:<20} {first_name:<20} {last_name:<20}\n")


class SearchChannels(TelegramParser):
    def __init__(self, api_id, api_hash, keywords: list[str]):
        super().__init__(api_id, api_hash)
        self.keywords = keywords

    @exception_handler
    async def found_channels(self):
        with open('founded_channels.html', 'w', encoding='utf-8') as file:
            file.write("<html><head><title>Найденные каналы</title></head>")
            file.write("<body>")
            file.write("<h2>Найденные каналы: </h2>")
            file.write("<ul>")

            for keyword in self.keywords:
                result = await self.client(SearchRequest(q=keyword, limit=150))
                for channel in result.chats:
                    if getattr(channel, 'megagroup', False):
                        channel_link = f"https://t.me/{channel.username}"
                        file.write("<li>")
                        file.write(f"<a href='{channel_link}'>{channel.title}</a> ({channel.id})")
                        file.write("</li><br>")

            file.write("</ul></body></html>")

    @exception_handler
    async def found_channels_txt(self):
        with open('founded_channels.txt', 'w', encoding='utf-8') as file:
            file.write("Найденные каналы:\n\n")

            for keyword in self.keywords:
                result = await self.client(SearchRequest(q=keyword, limit=150))
                for channel in result.chats:
                    if getattr(channel, 'megagroup', False):
                        channel_link = f"https://t.me/{channel.username}"
                        file.write(f"{channel.title} ({channel.id})\n")
                        file.write(f"Ссылка: {channel_link}\n\n")


async def main():
    api_id: int = 27523286
    api_hash: str = '2775f9dcb77f6ed0f86e327165192d56'
    channel: str = 'https://t.me/OstashkoNews'
    my_channel: str = 'https://t.me/coderchannel1'
    chat: str = 'https://t.me/leaders_talk'
    chat_id = 1253950227
    channel_id = -1001754252633
    keywords: list[str] = [
        "Автомобильные новости",
        "Автомобильные обзоры",
        "Тюнинг автомобилей",
        "Автоспорт",
        "Автомобильные технологии",
        "Автозапчасти",
        "Дорожные происшествия",
        "Автомобильные выставки",
        "Автомобильные аксессуары",
        "Электромобили"
    ]

    send_my_channel = TelegramParserChannel(api_id, api_hash, channel, my_channel)
    parser_chat = TelegramParserChat(api_id, api_hash, chat, chat_id)
    users_in_chat = TelegramParserChat(api_id, api_hash, chat)
    me = ShowMyChannels(api_id, api_hash)
    search_channels = SearchChannels(api_id, api_hash, keywords)

    async with me.client:
        await me.get_my_channels()


if __name__ == "__main__":
    import asyncio

    asyncio.run(main())
