import asyncio
import webbrowser
from abc import ABC, abstractmethod
from asyncio import sleep

from telethon import TelegramClient, types, events
from docx import Document
from telethon.tl.custom import Message
from telethon.tl.functions.contacts import SearchRequest

from telethon.tl.types import MessageMediaWebPage, MessageMediaPhoto, MessageMediaDocument, DocumentAttributeVideo

api_id = ''
api_hash = ''
client: TelegramClient = TelegramClient('anon', api_id, api_hash)

channels = 'https://t.me/RVvoenkor'
my_channel = 'https://t.me/coderchannel1'
discussion_channel_chat = 'https://t.me/rvvoenkor2'
chat = 'https://t.me/joinchat/yHTCCs3PpgkyMzIy'


class TelegramParser(ABC):
    def __init__(self, api_id, api_hash):
        self._api_id = api_id
        self._api_hash = api_hash
        self._client = client

    @abstractmethod
    async def parser_start(self) -> None:
        pass

    @property
    def client(self) -> TelegramClient:
        return self._client


class TelegramParserChannel(TelegramParser):
    def __init__(self, api_id: int, api_hash: str, channels=None, my_channel: str = None):
        super().__init__(api_id, api_hash)
        self.channels = channels
        self.my_channel = my_channel

    async def parser_start(self) -> None:
        document = Document()
        async for message in self.client.iter_messages(self.channels, limit=6):
            document.add_paragraph(f"ID: {message.id}")
            document.add_paragraph(f"Дата публикации: {message.date.strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph(f"Текст: {message.text}")
            document.add_paragraph("-" * 20)

        document.save("posts.docx")

    async def send_posts_to_my_channel(self):
        async for message in self.client.iter_messages(self.channels, limit=6):
            if message.text or message.media:
                if message.text and message.media:
                    await client.send_message(self.my_channel, message.text, file=message.media)
                elif message.text:
                    await client.send_message(self.my_channel, message.text)
                else:
                    await client.send_message(self.my_channel, file=message.media)


class TelegramParserChat(TelegramParser):
    def __init__(self, api_id, api_hash, chat):
        super().__init__(api_id, api_hash)
        self.chat = chat

    async def parser_start(self) -> None:
        document = Document()
        async for message in client.iter_messages(self.chat, limit=10):
            if message.text and message.sender:
                username = message.sender.username if message.sender.username else message.sender.first_name
                document.add_paragraph(f'Username: {username}')
                document.add_paragraph(f'Текст комментария: {message.text}')

        document.save('messages.docx')

    async def get_messages_html(self) -> None:
        html_content = "<html><head><title>Messages</title></head><body>"
        async for message in client.iter_messages(self.chat, limit=10):
            if message.text and message.sender:
                user_link = f"https://t.me/{message.sender.username}"

                html_content += f"<p>Username: <a href='{user_link}'>{message.sender.username}</a></p>"
                html_content += f"<p>Имя: {message.sender.first_name}</p>"
                html_content += f"<p>Фамилия: {message.sender.last_name}</p>"
                html_content += f"<p>Текст комментария: {message.text}</p>"
                html_content += "<hr>"

        html_content += "</body></html>"

        with open('messages.html', 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

    async def get_participants_chat(self):

        count: int = 0
        document = Document()
        async for participant in client.iter_participants(chat, limit=30):
            profile_link = f"https://t.me/{participant.username}"
            document.add_paragraph(f"User ID: {participant.id}")
            document.add_paragraph(f"Username: {profile_link}")
            document.add_paragraph(f"Имя: {participant.first_name}")
            document.add_paragraph(f"Фамилия: {participant.last_name}")
            document.add_paragraph(f"Телефон: {participant.phone}")
            document.add_paragraph(f"Имеет ли аккаунт премиум статус: {participant.premium}")
            document.add_paragraph("-" * 20)
            count += 1
        document.save('users.docx')
        print(f"Спарсено участников: {count}")

    async def get_participants_chat_html(self):
        count = 0
        html_content = "<html><head><title>Participants</title></head><body>"
        async for participant in client.iter_participants(chat, limit=30, filter=types.ChannelParticipantsAdmins):
            profile_link = f"https://t.me/{participant.username}"
            html_content += f"<p>User ID: <a href='{profile_link}'>{participant.id}</a></p>"
            html_content += f"<p>Username: <a href='{profile_link}'>{participant.username if participant.username else 'No username'}</a></p>"
            html_content += f"<p>Имя: {participant.first_name}</p>"
            html_content += f"<p>Фамилия: {participant.last_name}</p>"
            html_content += f"<p>Телефон: {participant.phone}</p>"
            html_content += f"<p>Имеет ли аккаунт премиум статус: {participant.premium}</p>"
            html_content += "<hr>"
            count += 1
        html_content += "</body></html>"

        with open('users.html', 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

        print(f"Спарсено участников: {count}")


class SearchChannels(TelegramParser):
    def __init__(self, api_id, api_hash, keywords: list[str]):
        super().__init__(api_id, api_hash)
        self.keywords = keywords

    async def parser_start(self):
        found_channels = []
        for keyword in self.keywords:
            result = await client(SearchRequest(q=keyword, limit=10))
            for channel in result.chats:
                if getattr(channel, 'megagroup', False):
                    channel_link = f"https://t.me/{channel.username}" if channel.username else f"https://t.me/c/{channel.id}"
                    found_channels.append(f"{channel.id} - {channel.title} - {channel_link}")

        print(' '.join(found_channels))


async def main():
    send_my_channel = TelegramParserChannel(api_id, api_hash, channels, my_channel)
    parser_chat = TelegramParserChat(api_id, api_hash, chat)
    users_in_chat = TelegramParserChat(api_id, api_hash, chat)
    search_channels = SearchChannels(api_id, api_hash, ["Bitcoin", "Ethereum", "crypto", "криптовалюта обучение"])

    async with client:
        users_in_chat = client.loop.create_task(search_channels.parser_start())

        await users_in_chat


if __name__ == "__main__":
    client.loop.run_until_complete(main())
